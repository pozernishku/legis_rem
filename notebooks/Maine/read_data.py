import io
from subprocess import Popen, PIPE
from docx import Document
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage

def convert_pdf_to_txt(path):
    rsrcmgr = PDFResourceManager()
    retstr = io.StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    # fp = open(path, 'rb')
    fp = path # path is _io.BufferedReader
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos = set()

    pdf_pages = PDFPage.get_pages(fp, pagenos, maxpages=maxpages,
                                  password=password,
                                  caching=caching,
                                  check_extractable=True)

    for page in pdf_pages:
        interpreter.process_page(page)

    text = retstr.getvalue()

    # fp.close()
    device.close()
    retstr.close()
    return text

def document_to_text(file_path):
    if file_path[-4:] == ".doc":
        cmd = ['antiword', file_path]
        p = Popen(cmd, stdout=PIPE)
        stdout, stderr = p.communicate()
        return stdout.decode('ascii', 'ignore')
    elif file_path[-5:] == ".docx":
        document = Document(file_path)
        paratextlist = document.paragraphs
        
        newparatextlist = []
        for paratext in paratextlist:
            newparatextlist.append(paratext.text)

        newtableparatextlist = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        newtableparatextlist.append(paragraph.text)
    
        return '\n\n'.join(newparatextlist + newtableparatextlist)
    elif file_path[-4:] == ".odt":
        cmd = ['odt2txt', file_path]
        p = Popen(cmd, stdout=PIPE)
        stdout, stderr = p.communicate()
        return stdout.decode('ascii', 'ignore')
